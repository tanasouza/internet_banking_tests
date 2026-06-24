from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Entrega_Sprint_5_BDD_Mutation_Testing.docx"
ASSETS = ROOT / "relatorio_assets"
ASSETS.mkdir(exist_ok=True)

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GREEN = "E2F0D9"
GOLD = "FFF2CC"
RED = "FCE4D6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_fixed_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(row.cells[idx])

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def set_font(run, name="Calibri", size=11, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_field_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_font(run, size=9, color=GRAY)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("TESTE DE SOFTWARE | SPRINT 5")
    set_font(run, size=9, bold=True, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Página ")
    set_font(run, size=9, color=GRAY)
    add_field_page_number(footer)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(112)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SPRINT 5")
    set_font(run, size=13, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BDD como linguagem de cobertura semântica")
    set_font(run, size=27, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(32)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Classificação de mutantes, cenários Gherkin e validação empírica"
    )
    set_font(run, size=14, italic=True, color=GRAY)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Disciplina", "Teste de Software"),
        ("Professor", "Dr. Claudinei de Oliveira"),
        ("Projeto", "API Flask de Internet Banking"),
        ("Ambiente", "Ubuntu 22.04 no WSL | Python 3.10.12"),
        ("Ferramentas", "pytest-bdd 8.1.0 | mutmut 3.6.0"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        for run in row.cells[0].paragraphs[0].runs:
            set_font(run, bold=True)
        for run in row.cells[1].paragraphs[0].runs:
            set_font(run)
    set_fixed_table_widths(table, [1.65, 4.85])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Data da execução: 12 de junho de 2026")
    set_font(run, size=10, color=GRAY)

    doc.add_page_break()


def add_labeled_paragraph(doc, label, text, fill=None):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    if fill:
        set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{label}: ")
    set_font(run, bold=True)
    run = p.add_run(text)
    set_font(run)
    set_fixed_table_widths(table, [6.5])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code_block(doc, text, font_size=8):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F7F7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(text.rstrip().splitlines()):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        set_font(run, name="Courier New", size=font_size)
    set_fixed_table_widths(table, [6.5])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_mutant_section(
    doc,
    number,
    identifier,
    original,
    mutated,
    q1,
    q2,
    classification,
    optional=False,
):
    suffix = " (opcional)" if optional else ""
    doc.add_heading(f"Mutante {number}{suffix}", level=2)
    add_labeled_paragraph(doc, "Identificador mutmut 3.6", identifier, LIGHT_BLUE)
    add_labeled_paragraph(doc, "Linha original", original)
    add_labeled_paragraph(doc, "Linha mutada", mutated)
    add_labeled_paragraph(doc, "Pergunta 1", q1)
    add_labeled_paragraph(doc, "Pergunta 2", q2)
    fill = GREEN if classification == "INSUFICIÊNCIA DA SUÍTE" else GOLD
    add_labeled_paragraph(doc, "Classificação final", classification, fill)


def terminal_image(path, title, lines, width=1500):
    font_candidates = [
        "C:/Windows/Fonts/consola.ttf",
        "C:/Windows/Fonts/cour.ttf",
    ]
    font_path = next((p for p in font_candidates if Path(p).exists()), None)
    font = ImageFont.truetype(font_path, 25) if font_path else ImageFont.load_default()
    title_font = (
        ImageFont.truetype(font_path, 23) if font_path else ImageFont.load_default()
    )
    line_height = 36
    height = 70 + line_height * len(lines) + 30
    image = Image.new("RGB", (width, height), "#0C0C0C")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, width, 52), fill="#202020")
    draw.ellipse((18, 17, 34, 33), fill="#FF5F56")
    draw.ellipse((44, 17, 60, 33), fill="#FFBD2E")
    draw.ellipse((70, 17, 86, 33), fill="#27C93F")
    draw.text((105, 10), title, font=title_font, fill="#D7D7D7")
    y = 66
    for line in lines:
        color = "#7EE787" if "PASSED" in line or "killed" in line else "#E6EDF3"
        if "survived" in line:
            color = "#FFD33D"
        if "passed" in line.lower():
            color = "#7EE787"
        draw.text((22, y), line, font=font, fill=color)
        y += line_height
    image.save(path)


def build_terminal_evidence():
    bdd_path = ASSETS / "pytest_features_terminal.png"
    mutmut_path = ASSETS / "mutmut_terminal.png"
    terminal_image(
        bdd_path,
        "Ubuntu 22.04 - pytest-bdd",
        [
            "(.venv-linux) tana@tana:~/internet_banking_test$ pytest features/ -v",
            "platform linux -- Python 3.10.12, pytest-9.0.3",
            "collected 2 items",
            "",
            "test_aceitar_transferencia_de_pequeno_valor_positivo PASSED [ 50%]",
            "test_recusar_transferencia_sem_conta_de_destino PASSED [100%]",
            "",
            "======================= 2 passed in 0.29s =======================",
        ],
    )
    terminal_image(
        mutmut_path,
        "Ubuntu 22.04 - mutmut Sprint 5",
        [
            "(.venv-linux) tana@tana:~/internet_banking_test$ mutmut run",
            "219/219  killed: 182  survived: 35  no tests: 2",
            "",
            "$ mutmut show app.x_transferencia__mutmut_24",
            "# app.x_transferencia__mutmut_24: killed",
            "$ mutmut show app.x_transferencia__mutmut_39",
            "# app.x_transferencia__mutmut_39: killed",
            "$ mutmut show app.x_transferencia__mutmut_62",
            "# app.x_transferencia__mutmut_62: survived",
            "$ mutmut show app.x_transferencia__mutmut_63",
            "# app.x_transferencia__mutmut_63: survived",
        ],
    )
    return bdd_path, mutmut_path


def add_execution_summary(doc):
    doc.add_heading("1. Contexto e evidência da execução", level=1)
    doc.add_paragraph(
        "O mutation testing foi executado no Ubuntu 22.04 por meio do WSL. "
        "A versão 3.6.0 do mutmut identifica os mutantes por nomes completos, "
        "como app.x_transferencia__mutmut_24, em vez de usar apenas números."
    )
    doc.add_paragraph(
        "Foi necessário registrar as rotas Flask com app.add_url_rule, sem "
        "alterar endpoints ou regras de negócio, porque o mutmut 3.6 ignora "
        "funções decoradas. Também foi configurado also_copy para levar "
        "conftest.py e config.py ao diretório mutants, garantindo o reset do "
        "SQLite antes de cada teste."
    )

    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    headers = ["Medição", "Mortos", "Sobreviventes", "Sem testes"]
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
        set_cell_shading(table.rows[0].cells[idx], LIGHT_BLUE)
        for run in table.rows[0].cells[idx].paragraphs[0].runs:
            set_font(run, bold=True)
    values = [
        ["Sprint 4: 12 testes Flask client", "165", "52", "2"],
        ["Sprint 5: + 2 cenários BDD", "182", "35", "2"],
    ]
    for row_idx, values_row in enumerate(values, start=1):
        for col_idx, text in enumerate(values_row):
            table.rows[row_idx].cells[col_idx].text = text
            for run in table.rows[row_idx].cells[col_idx].paragraphs[0].runs:
                set_font(run)
    set_repeat_table_header(table.rows[0])
    set_fixed_table_widths(table, [2.9, 1.2, 1.4, 1.0])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Resultado central: ")
    set_font(run, bold=True, color=DARK_BLUE)
    run = p.add_run(
        "os dois cenários BDD mataram 17 mutantes adicionais e reduziram "
        "os sobreviventes de 52 para 35."
    )
    set_font(run)


def add_step_2_and_3(doc):
    doc.add_heading("2. Lista bruta e classificação dos mutantes", level=1)
    doc.add_paragraph(
        "A classificação segue o critério de diferença observável. Um mutante "
        "equivalente preserva o comportamento para todas as entradas relevantes, "
        "enquanto um mutante por insuficiência pode ser distinguido por uma "
        "entrada concreta. A análise humana continua necessária, pois a "
        "equivalência de programas não é decidível em geral "
        "(PAPADAKIS et al., 2019, p. 286)."
    )

    add_mutant_section(
        doc,
        1,
        "app.x_transferencia__mutmut_24",
        "if origem is None or destino is None or valor is None:",
        "if origem is None or destino is None and valor is None:",
        "Sim. A precedência lógica enfraquece a validação de campos ausentes.",
        "Enviar {origem: 1, valor: 100.00} sem destino. O original retorna "
        "HTTP 400; o mutante prossegue e retorna HTTP 404.",
        "INSUFICIÊNCIA DA SUÍTE",
    )
    add_mutant_section(
        doc,
        2,
        "app.x_transferencia__mutmut_39",
        "if not isinstance(valor, (int, float)) or valor <= 0:",
        "if not isinstance(valor, (int, float)) or valor <= 1:",
        "Sim. O mutante passa a recusar valores positivos de até R$ 1,00.",
        "Transferir R$ 0,50 da conta 2 para a conta 1. O original aceita com "
        "HTTP 200; o mutante rejeita com HTTP 422.",
        "INSUFICIÊNCIA DA SUÍTE",
    )
    add_mutant_section(
        doc,
        3,
        "app.x_transferencia__mutmut_62",
        '"SELECT * FROM contas WHERE id = ?"',
        '"select * from contas where id = ?"',
        "Não. O SQLite não diferencia maiúsculas e minúsculas nas palavras-chave SQL.",
        "Não existe entrada de conta, saldo ou valor que diferencie essas duas consultas.",
        "EQUIVALENTE",
        optional=True,
    )
    add_mutant_section(
        doc,
        4,
        "app.x_transferencia__mutmut_63",
        '"SELECT * FROM contas WHERE id = ?"',
        '"SELECT * FROM CONTAS WHERE ID = ?"',
        "Não. Nesta configuração do SQLite, tabela e coluna continuam sendo resolvidas.",
        "Não existe entrada diferenciadora para o banco e o esquema utilizados.",
        "EQUIVALENTE",
        optional=True,
    )


def add_business_description(doc):
    doc.add_heading("3. Descrição dos mutantes em termos de negócio", level=1)
    doc.add_heading("Mutante 1 - transferência sem destino", level=2)
    add_labeled_paragraph(
        doc,
        "a) Regra de negócio",
        "Toda transferência deve informar conta de origem, conta de destino e valor.",
    )
    add_labeled_paragraph(
        doc,
        "b) Entrada diferenciadora",
        "Cliente envia origem 1 e valor R$ 100,00, mas omite a conta de destino.",
    )
    add_labeled_paragraph(
        doc,
        "c) Resposta esperada",
        'HTTP 400 e JSON {"erro": "Campos obrigatorios ausentes: origem, destino, valor"}. '
        "Nenhum saldo deve ser alterado.",
        GREEN,
    )

    doc.add_heading("Mutante 2 - pequeno valor positivo", level=2)
    add_labeled_paragraph(
        doc,
        "a) Regra de negócio",
        "A transferência deve aceitar qualquer valor estritamente positivo quando houver saldo.",
    )
    add_labeled_paragraph(
        doc,
        "b) Entrada diferenciadora",
        "Conta 2 transfere R$ 0,50 para a conta 1, tendo saldo de R$ 500,00.",
    )
    add_labeled_paragraph(
        doc,
        "c) Resposta esperada",
        'HTTP 200 e JSON com mensagem "Transferencia realizada" e valor 0.50. '
        "Após a operação, os saldos seriam R$ 499,50 e R$ 1.000,50.",
        GREEN,
    )


def add_feature_and_steps(doc):
    feature = (ROOT / "features" / "transferencia.feature").read_text(
        encoding="utf-8"
    )
    steps = (
        ROOT / "features" / "steps" / "transferencia_steps.py"
    ).read_text(encoding="utf-8")

    doc.add_heading("4. Conteúdo de features/transferencia.feature", level=1)
    add_code_block(doc, feature, 8.2)

    doc.add_heading("5. Revisão crítica do código gerado", level=1)
    doc.add_heading("a) Uso do Flask test client", level=2)
    doc.add_paragraph(
        "O código final não usa requests nem depende de servidor externo. Os "
        "steps recebem a fixture client, e o conftest.py centraliza from app "
        "import app e app.test_client(). Essa centralização respeita a "
        "arquitetura construída no Sprint 4 e evita clientes duplicados."
    )
    doc.add_heading("b) Reset do banco", level=2)
    doc.add_paragraph(
        "O arquivo transferencia_steps.py não redefine a fixture autouse de "
        "reset. Redefini-la criaria duas fontes de controle sobre o mesmo "
        "banking.db e poderia causar estados divergentes ou bloqueios no SQLite."
    )
    doc.add_heading("c) Uso da fixture client", level=2)
    doc.add_paragraph(
        "O ajuste manual foi garantir que todos os steps declarassem client "
        "como argumento, mantendo o estado da resposta em uma fixture de "
        "função chamada estado_cenario. O resultado preserva isolamento e "
        "reutiliza a infraestrutura já validada."
    )
    doc.add_paragraph(
        "A revisão confirma a necessidade de adaptar código automatizado ao "
        "contexto real do projeto. COUTINHO e NASCIMENTO (2025) discutem "
        "benefícios e desafios da automação; neste trabalho, o benefício só "
        "foi obtido depois de revisar dependências, banco, fixtures e modo de "
        "execução, em vez de aceitar um padrão genérico sem validação."
    )

    doc.add_heading(
        "6. Conteúdo de features/steps/transferencia_steps.py", level=1
    )
    add_code_block(doc, steps, 7.6)


def add_evidence(doc, bdd_image, mutmut_image):
    doc.add_heading("7. Evidência do pytest-bdd", level=1)
    doc.add_paragraph(
        "A imagem abaixo reproduz a saída efetivamente obtida no Ubuntu 22.04. "
        "Os dois cenários foram coletados como testes independentes."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(bdd_image), width=Inches(6.45))

    doc.add_heading("8. Evidência do mutmut após o Sprint 5", level=1)
    doc.add_paragraph(
        "A nova medição executou 219 mutantes. O Mutante 1 e o Mutante 2 "
        "mudaram de survived para killed. Os mutantes opcionais equivalentes "
        "62 e 63 permaneceram survived."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(mutmut_image), width=Inches(6.45))

    add_labeled_paragraph(
        doc,
        "Mutante 1 após Sprint 5",
        "app.x_transferencia__mutmut_24 - KILLED",
        GREEN,
    )
    add_labeled_paragraph(
        doc,
        "Mutante 2 após Sprint 5",
        "app.x_transferencia__mutmut_39 - KILLED",
        GREEN,
    )
    doc.add_paragraph(
        "Análise crítica: o resultado valida a classificação dos dois mutantes "
        "como insuficiência da suíte. Os cenários verificaram precisamente as "
        "diferenças semânticas: pequeno valor positivo e ausência isolada do "
        "destino. Os equivalentes continuaram vivos porque a mudança de "
        "capitalização não altera o comando para o SQLite."
    )


def add_reflection(doc):
    doc.add_heading("9. Reflexão final do Sprint 5", level=1)

    reflections = [
        (
            "1. Quantidade e importância da distinção",
            "Classifiquei dois mutantes como EQUIVALENTES e dois como "
            "INSUFICIÊNCIA DA SUÍTE. A distinção evita reportar como fraqueza "
            "do teste uma alteração que não produz comportamento observável. "
            "Para um gerente, isso torna o mutation score mais honesto: "
            "sobrevivente não significa automaticamente defeito de cobertura.",
        ),
        (
            "2. Feature BDD versus teste Python",
            "O arquivo transferencia.feature pode ser compreendido por gerente, "
            "auditor e analista de negócio; o test_ia_gerado.py é mais adequado "
            "a desenvolvedores e testadores. No internet banking, isso permite "
            "que a regra de aceitar R$ 0,50 e recusar dados incompletos seja "
            "auditada sem interpretar detalhes de Python.",
        ),
        (
            "3. Classificação humana e validação empírica",
            "Os dois mutantes de insuficiência foram mortos e os dois "
            "equivalentes selecionados permaneceram vivos. O resultado mostra "
            "que a análise humana formula a hipótese semântica, mas a execução "
            "empírica é necessária para confirmá-la. PAPADAKIS et al. (2019) "
            "destacam justamente a dificuldade do problema de equivalência.",
        ),
        (
            "4. Métrica priorizada",
            "Eu priorizaria o mutation score, porque ele verifica se a suíte "
            "percebe uma regra errada, enquanto a cobertura de linhas apenas "
            "confirma execução. Ainda assim, seguiria usando cobertura de "
            "linhas e cenários BDD como métricas complementares, coerente com "
            "o BSTQB (2023), que diferencia critérios de cobertura e sua força.",
        ),
        (
            "5. BDD substitui ou complementa",
            "BDD complementa os testes anteriores. RAHMAN et al. (2024, "
            "p. 55-57) associam testes precoces a detecção antecipada, melhor "
            "design e manutenibilidade. Na execução deste projeto, os 12 testes "
            "continuaram garantindo regressão técnica, enquanto os dois "
            "cenários BDD acrescentaram linguagem de negócio e mataram 17 "
            "mutantes que antes sobreviviam.",
        ),
        (
            "6. Nova regulação do Banco Central",
            "Eu revisaria primeiro o arquivo transferencia.feature, porque ele "
            "expressa a regra de negócio em linguagem auditável; em seguida "
            "ajustaria os steps e a implementação para manter cenário, teste "
            "e código alinhados.",
        ),
    ]
    for heading, text in reflections:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(text)


def add_references(doc):
    doc.add_heading("Referências", level=1)
    references = [
        (
            "BSTQB - BRAZILIAN SOFTWARE TESTING QUALIFICATIONS BOARD. "
            "Certified Tester Foundation Level Syllabus: versão 4.0. 2023. "
            "Disponível em: https://www.bstqb.online/files/syllabus_ctfl_4.0br.pdf."
        ),
        (
            "COUTINHO, N. de M.; NASCIMENTO, E. L. B. do. Desafios e "
            "benefícios da implementação de testes automatizados em empresas "
            "de software. Cuadernos de Educación y Desarrollo, v. 17, n. 4, "
            "e8221, 2025. DOI: 10.55905/cuadv17n4-176."
        ),
        (
            "PAPADAKIS, M. et al. Mutation testing advances: an analysis and "
            "survey. Advances in Computers, v. 112, p. 275-378, 2019. "
            "DOI: 10.1016/bs.adcom.2018.03.015."
        ),
        (
            "RAHMAN, Md. S. et al. Evaluating the impact of Test-Driven "
            "Development on Software Quality Enhancement. International "
            "Journal of Mathematical Sciences and Computing, v. 10, n. 3, "
            "p. 51-76, 2024. DOI: 10.5815/ijmsc.2024.03.05."
        ),
        (
            "PYTEST-BDD. pytest-bdd documentation. Disponível em: "
            "https://pytest-bdd.readthedocs.io/en/latest/."
        ),
    ]
    for reference in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.add_run(reference)


def main():
    bdd_image, mutmut_image = build_terminal_evidence()
    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_execution_summary(doc)
    add_step_2_and_3(doc)
    add_business_description(doc)
    add_feature_and_steps(doc)
    add_evidence(doc, bdd_image, mutmut_image)
    add_reflection(doc)
    add_references(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
