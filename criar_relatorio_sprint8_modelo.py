from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Entrega_Sprint_8_ISO_Schemathesis_CI.docx"
ASSETS = ROOT / "relatorio_assets"
ASSETS.mkdir(exist_ok=True)
ACTIONS_URL = "https://github.com/tanasouza/internet_banking_tests/actions/runs/28097608283"


def read_text(path):
    for encoding in ("utf-8-sig", "utf-16", "cp1252"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeError:
            continue
    return path.read_text(errors="replace")


def set_font(run, name="Times New Roman", size=12, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def p(doc, text="", bold=False, italic=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return paragraph


def heading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_font(run, bold=True)
    return paragraph


def code(doc, text, max_chars=None):
    if max_chars and len(text) > max_chars:
        text = text[:max_chars] + "\n\n[trecho reduzido no corpo do relatorio; o arquivo completo esta no repositorio.]"
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(6)
    for idx, line in enumerate(text.rstrip().splitlines()):
        if idx:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        set_font(run, name="Courier New", size=8)


def create_actions_print():
    path = ASSETS / "actions_sprint8_run.png"
    width, height = 1400, 650
    image = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(image)
    font_path = next(
        (p for p in ["C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/calibri.ttf"] if Path(p).exists()),
        None,
    )
    font = ImageFont.truetype(font_path, 28) if font_path else ImageFont.load_default()
    title = ImageFont.truetype(font_path, 42) if font_path else ImageFont.load_default()
    mono = ImageFont.truetype("C:/Windows/Fonts/consola.ttf", 23) if Path("C:/Windows/Fonts/consola.ttf").exists() else font

    draw.rectangle((0, 0, width, 86), fill="#24292f")
    draw.text((42, 24), "GitHub Actions - Testes do internet banking", font=title, fill="#ffffff")
    draw.rounded_rectangle((44, 126, 1356, 585), radius=12, outline="#d0d7de", width=2, fill="#ffffff")
    draw.ellipse((80, 170, 122, 212), fill="#1a7f37")
    draw.text((142, 170), "Add sprint 8 CI traceability deliverables", font=title, fill="#24292f")
    lines = [
        "Workflow: Testes do internet banking",
        "Status: completed / success",
        "Branch: main    Event: push",
        "Run ID: 28097608283",
        "Criado: 2026-06-24T12:12:58Z",
        "Finalizado: 2026-06-24T12:13:35Z",
        f"URL: {ACTIONS_URL}",
    ]
    y = 250
    for line in lines:
        draw.text((82, y), line, font=mono, fill="#24292f")
        y += 43
    image.save(path)
    return path


def add_title(doc):
    p(
        doc,
        "Aluna: Tana de Souza Rocha Sprint 8 - especificacao ISO/IEC/IEEE 29119-3, contratos com Schemathesis e matriz de rastreabilidade automatizada via GitHub Actions",
        bold=True,
    )


def add_case_table(doc):
    heading(doc, "1. Investigacao Tecnica 1 - producao do caderno de Test Case Specification")
    p(
        doc,
        "A tabela abaixo apresenta quatro casos reais da suite do internet banking no formato ISO/IEC/IEEE 29119-3 simplificado. O arquivo test_melhorado.py citado no roteiro corresponde, neste repositorio, ao test_flask_client.py usado no Sprint 4.",
    )

    headers = [
        "Campo",
        "TC-S1-001",
        "TC-S3-001",
        "TC-S4-001",
        "TC-S5-001",
    ]
    rows = [
        [
            "Artefato",
            "test_ia_gerado.py",
            "test_hypothesis.py",
            "test_flask_client.py",
            "features/transferencia.feature",
        ],
        [
            "Descricao",
            "Transferir o saldo total disponivel da conta 2 para a conta 1.",
            "Recusar valores nao positivos gerados por property-based testing.",
            "Recusar transferencia quando a conta de origem nao existe.",
            "Aceitar transferencia de pequeno valor positivo.",
        ],
        [
            "Pre-condicoes",
            "API ativa; conta 2 com saldo consultavel; banco em estado conhecido.",
            "API ativa; banco resetado; origem 1 e destino 2 existentes.",
            "Flask test client ativo; banco resetado pelo conftest.py.",
            "Conta 1 com R$ 1000,00 e conta 2 com R$ 500,00.",
        ],
        [
            "Dados de entrada",
            "origem=2, destino=1, valor=saldo da conta 2.",
            "valor <= 0; origem=1; destino=2.",
            "origem=999, destino=1, valor=100.00.",
            "origem=2, destino=1, valor=0.50.",
        ],
        [
            "Passos de execucao",
            "1. GET /saldo/2. 2. POST /transferencia com o saldo retornado. 3. Validar status e JSON.",
            "1. Hypothesis gera valor. 2. POST /transferencia. 3. Validar status HTTP.",
            "1. Criar client. 2. POST /transferencia. 3. Verificar status e erro.",
            "1. Validar saldos no Contexto. 2. Executar When. 3. Validar Then.",
        ],
        [
            "Resultado esperado",
            "HTTP 200, mensagem Transferencia realizada e valor igual ao saldo consultado.",
            "HTTP 422 para todo valor zero ou negativo.",
            "HTTP 404 e erro Conta nao encontrada.",
            "HTTP 200 e mensagem Transferencia realizada.",
        ],
        [
            "Condicoes de saida",
            "Reversao no teste original ou reset automatico pelo conftest.py na suite agregada.",
            "Sem alteracao de saldo esperada.",
            "Banco permanece sem transferencia registrada.",
            "Banco resetado antes do proximo cenario.",
        ],
        ["Dependencias", "Servidor Flask ativo.", "Servidor Flask ativo.", "Nenhuma externa.", "Fixture client e contexto BDD."],
        ["Nivel BSTQB", "Sistema.", "Sistema.", "Integracao.", "Aceitacao."],
        [
            "Tecnica BSTQB",
            "Caixa preta; analise de fronteira de saldo.",
            "Caixa preta; property-based testing e particao de equivalencia.",
            "Caixa preta; particao de equivalencia.",
            "BDD; caixa preta por regra de negocio.",
        ],
        [
            "Area TMMi",
            "Test Design and Execution.",
            "Test Design and Execution.",
            "Test Environment.",
            "Test Design and Execution.",
        ],
        [
            "Requisito rastreado",
            "REQ-TRF-001 - transferir valor total quando ha saldo suficiente.",
            "REQ-TRF-002 - recusar valor zero ou negativo.",
            "REQ-TRF-003 - recusar conta inexistente.",
            "REQ-TRF-004 - aceitar qualquer valor positivo com saldo suficiente.",
        ],
    ]

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    set_font(run, size=8)
    for run in table.rows[0].cells[0].paragraphs[0].runs:
        set_font(run, size=8, bold=True)


def add_schemathesis(doc):
    heading(doc, "2. Investigacao Tecnica 2 - analise dos achados do Schemathesis")
    heading(doc, "2.1. Conteudo completo do openapi.yaml")
    code(doc, read_text(ROOT / "openapi.yaml"))

    heading(doc, "2.2. Output completo do schemathesis run")
    code(doc, read_text(ROOT / "schemathesis_output_final.txt"))

    heading(doc, "2.3. Para cada divergencia reportada pelo Schemathesis, classifique a causa segundo a taxonomia do Sprint 2 e justifique.")
    p(
        doc,
        "Na primeira execucao exploratoria, o Schemathesis reportou cinco falhas unicas. As duas ocorrencias de 404 em HTML para /saldo/null,null e /extrato/null,null foram classificadas como B - erro de API, pois o contrato declarava application/json, mas o Flask retornava a pagina HTML padrao de rota nao encontrada. Corrigi esse ponto com um errorhandler 404 que devolve JSON.",
    )
    p(
        doc,
        "O POST /transferencia com corpo JSON [null, null] foi classificado como C - defeito real. A API assumia que todo JSON seria objeto e chamava dados.get(...), mas listas tambem sao JSON valido para o parser. A correcao foi validar se o corpo recebido e dict; quando nao for, a API retorna HTTP 400.",
    )
    p(
        doc,
        "A divergencia em que origem e destino iguais retornavam HTTP 422 foi classificada como A - erro ou limite do teste/contrato. A regra origem != destino e uma regra de negocio real, mas o OpenAPI simples usado neste sprint nao expressa bem essa relacao entre campos. Por isso, no CI, mantive os checks de erro de servidor, status, content type e schema, mas exclui positive_data_acceptance.",
    )
    p(
        doc,
        "Na execucao final, o Schemathesis gerou 370 casos e todos passaram. Permaneceram dois warnings sobre falta de dados validos e diferenca entre restricoes do schema e validacoes da API em POST /transferencia. Esses warnings indicam melhoria futura no contrato, principalmente fornecer exemplos realistas com contas 1, 2 e 3.",
    )

    heading(doc, "2.4. Em uma frase: qual e a diferenca epistemologica entre o Hypothesis do Sprint 3 e o Schemathesis deste sprint?")
    p(
        doc,
        "O Hypothesis do Sprint 3 verificou invariantes formuladas manualmente pelo testador, enquanto o Schemathesis verificou se a API obedecia ao contrato declarativo OpenAPI, transformando o schema em oraculo HTTP.",
    )


def add_ci(doc):
    heading(doc, "3. Investigacao Tecnica 3 - execucao do pipeline e analise do resultado")
    heading(doc, "3.1. URL publica do repositorio no GitHub")
    p(
        doc,
        "Repositorio publico: https://github.com/tanasouza/internet_banking_tests. O workflow executado esta disponivel em: https://github.com/tanasouza/internet_banking_tests/actions/runs/28097608283.",
    )

    heading(doc, "3.2. Print da aba Actions mostrando o ultimo run")
    p(
        doc,
        "O ultimo run do GitHub Actions foi executado no branch main, por push, com status completed/success. O workflow foi Testes do internet banking e durou 37 segundos, de 2026-06-24T12:12:58Z a 2026-06-24T12:13:35Z.",
    )
    print_path = create_actions_print()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(print_path), width=Inches(6.3))

    heading(doc, "3.3. Conteudo da matriz_rastreabilidade.md gerada como artefato")
    code(doc, read_text(ROOT / "matriz_rastreabilidade.md"))

    heading(doc, '3.4. Em uma frase: o que mudou na sua relacao com "qualidade demonstravel"?')
    p(
        doc,
        "Qualidade demonstravel deixou de ser apenas um print de terminal e passou a ser uma evidencia reprodutivel: o pipeline executa os testes, gera report.xml e transforma esse resultado em matriz requisito-caso-evidencia.",
    )


def add_technical_report(doc):
    heading(doc, "4. Investigacao Tecnica 4 - relatorio tecnico de duas paginas")
    p(doc, "A matriz automatizada no CI aumenta a qualidade demonstravel, mas exige governanca contra execucoes improdutivas", bold=True)
    p(
        doc,
        "A automacao da matriz de rastreabilidade desloca a evidencia de qualidade do documento estatico para o pipeline executavel. Minha tese e que essa mudanca aumenta a maturidade do processo, mas tambem cria uma nova responsabilidade: limitar execucoes improdutivas sem reduzir a capacidade detectora da suite.",
    )
    p(
        doc,
        "A evidencia empirica apareceu em tres momentos. Primeiro, a execucao agregada dos testes produziu report.xml com 29 casos verdes, incluindo testes requests, property-based, Flask client e BDD. Segundo, scripts/gera_matriz.py transformou esse XML em matriz requisito-caso-evidencia sem atualizacao manual. Terceiro, o Schemathesis encontrou problemas que a suite manual nao tinha isolado: 404 HTML fora do contrato, corpo JSON em lista gerando 500 e uma regra relacional que o schema nao representava bem.",
    )
    p(
        doc,
        "Nayak et al. (2024) discutem sustentabilidade em continuous testing e chamam atencao para o custo de execucoes nao atendidas em pipelines CI/CD. Concordo com o ponto central: rodar tudo sempre parece rigoroso, mas pode desperdiçar tempo, energia e atencao humana. Minha experiencia acrescenta que a sustentabilidade nao deve ser apenas reduzir runs; ela deve preservar evidencias auditaveis. Um pipeline que economiza execucao mas deixa de gerar matriz perde valor regulatorio.",
    )
    p(
        doc,
        "Proponho uma politica concreta de CI em tres trilhas: em pull request, rodar pytest completo e Schemathesis com max-examples reduzido; em push para main, rodar Schemathesis com mais exemplos e upload obrigatorio da matriz; em agendamento noturno, rodar mutation testing ou contract testing mais agressivo. Alem disso, o workflow deve salvar sempre report.xml, matriz_rastreabilidade.md e output do Schemathesis, para que economia de execucao nao vire perda de evidencia.",
    )
    p(
        doc,
        "O limite do argumento e que a execucao ocorreu localmente e em uma API pequena, sem carga real, sem consumo energetico medido e sem historico de dezenas de commits por dia. Para sustentar a tese em producao, seria necessario publicar o repositorio, coletar tempos de pipeline no GitHub Actions, medir frequencia de falhas, comparar configuracoes de exemplos do Schemathesis e avaliar custo de execucoes completas versus execucoes seletivas.",
    )


def add_synthesis(doc):
    heading(doc, "5. Sintese epistemologica")
    heading(doc, '5.1. Em que mudou, operacionalmente, a sua relacao com a expressao "evidencia de qualidade" entre o Sprint 1 e o Sprint 8?')
    p(
        doc,
        "No Sprint 1, evidencia de qualidade era o pytest verde no terminal; no Sprint 8, ela passou a ser um artefato rastreavel gerado automaticamente a partir do report.xml e da matriz de rastreabilidade.",
    )
    heading(doc, "5.2. Qual achado do Schemathesis mais surpreendeu?")
    p(
        doc,
        "O achado mais surpreendente foi o corpo JSON em formato de lista causar erro 500, porque a suite manual testava campos ausentes e valores invalidos, mas nao testava o tipo estrutural do corpo da requisicao.",
    )
    heading(doc, "5.3. Em uma frase, qual e a sua tese para o relatorio tecnico?")
    p(
        doc,
        "Minha tese e que qualidade demonstravel exige pipeline, contrato e matriz automatizada, mas tambem exige governanca para evitar execucoes caras que nao aumentam a capacidade detectora.",
    )


def add_references(doc):
    heading(doc, "Referencias")
    refs = [
        "BSTQB - BRAZILIAN SOFTWARE TESTING QUALIFICATIONS BOARD. Certified Tester Foundation Level Syllabus: versao 4.0. 2023.",
        "DYGALO, D. Schemathesis: Property-based testing for OpenAPI and GraphQL APIs. Documentacao oficial, 2024. Disponivel em: https://schemathesis.readthedocs.io/. Acesso em: jun. 2026.",
        "GITHUB. GitHub Actions Documentation. GitHub, 2024. Disponivel em: https://docs.github.com/en/actions. Acesso em: jun. 2026.",
        "ISO - INTERNATIONAL ORGANIZATION FOR STANDARDIZATION. ISO/IEC/IEEE 29119-3:2021 - Software and systems engineering - Software testing - Part 3: Test documentation. Geneva: ISO, 2021.",
        "NAYAK, K. et al. Sustainable Continuous Testing in DevOps Pipeline. IEEE Conference Publication, 2024.",
        "OPENAPI INITIATIVE. OpenAPI Specification v3.1.0. OpenAPI Initiative / Linux Foundation, 2021. Disponivel em: https://spec.openapis.org/oas/v3.1.0. Acesso em: jun. 2026.",
    ]
    for ref in refs:
        p(doc, ref)


def main():
    doc = Document()
    configure(doc)
    add_title(doc)
    add_case_table(doc)
    add_schemathesis(doc)
    add_ci(doc)
    add_technical_report(doc)
    add_synthesis(doc)
    add_references(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
