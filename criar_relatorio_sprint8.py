from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Entrega_Sprint_8_ISO_Schemathesis_CI.docx"

BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_GREEN = "E2F0D9"
LIGHT_GOLD = "FFF2CC"


def set_font(run, name="Times New Roman", size=12, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    set_font(run, size=10, color=MUTED)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, before, after in (
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("Teste de Software | Sprint 8")
    set_font(r, size=10, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Pagina ")
    set_font(r, size=10, color=MUTED)
    page_number(footer)


def para(doc, text="", prefix=None):
    p = doc.add_paragraph()
    if prefix:
        r = p.add_run(prefix)
        set_font(r, bold=True)
        r = p.add_run(text)
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def code_block(doc, text, size=7.2, max_chars=None):
    if max_chars and len(text) > max_chars:
        text = text[:max_chars] + "\n\n[trecho truncado no corpo do Word; arquivo completo esta no repositorio.]"
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, "F7F7F7")
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(text.rstrip().splitlines()):
        if idx:
            p.add_run("\n")
        r = p.add_run(line)
        set_font(r, name="Courier New", size=size)
    set_table_widths(table, [6.5])


def read_text(path):
    for encoding in ("utf-8-sig", "utf-16", "cp1252"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeError:
            continue
    return path.read_text(errors="replace")


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SPRINT 8")
    set_font(r, size=13, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ISO 29119-3, Schemathesis e Matriz de Rastreabilidade no CI")
    set_font(r, size=23, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Projeto: internet_banking_tests")
    set_font(r, size=13, italic=True, color=MUTED)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Disciplina", "Teste de Software"),
        ("Professor", "Dr. Claudinei de Oliveira"),
        ("Objeto", "API Flask de Internet Banking e suite dos Sprints 1 a 5"),
        ("Artefatos Sprint 8", "openapi.yaml, requirements.txt, .github/workflows/test.yml, scripts/gera_matriz.py, report.xml e matriz_rastreabilidade.md"),
        ("Data", "24 de junho de 2026"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        shade(row.cells[0], LIGHT_BLUE)
        for run in row.cells[0].paragraphs[0].runs:
            set_font(run, bold=True)
        for run in row.cells[1].paragraphs[0].runs:
            set_font(run)
    set_table_widths(table, [1.65, 4.85])
    doc.add_page_break()


def add_cases(doc):
    doc.add_heading("1. Caderno de Test Case Specification - ISO/IEC/IEEE 29119-3 simplificado", level=1)
    para(doc, "A tabela especifica quatro casos reais da suite, um por artefato principal dos sprints anteriores. O arquivo test_melhorado.py citado no roteiro corresponde, neste repositorio, ao test_flask_client.py produzido para o Sprint 4.")

    rows = [
        ("TC-S1-001", "Transferencia com saldo igual ao valor disponivel.", "Banco resetado; conta 2 com R$ 500,00 e conta 1 existente.", "origem=2, destino=1, valor=saldo da conta 2", "1. Consultar /saldo/2. 2. POST /transferencia com o saldo retornado. 3. Validar resposta.", "HTTP 200, mensagem Transferencia realizada, valor igual ao saldo consultado.", "Reversao por transferencia inversa no teste original; reset por conftest na suite agregada.", "Nenhuma.", "Sistema", "Caixa preta; analise de fronteira de saldo.", "Test Design and Execution", "REQ-TRF-001"),
        ("TC-S3-001", "Recusar transferencia com valor nao positivo gerado por propriedade.", "Servidor ativo; banco resetado por fixture autouse.", "valor em [-1000000, 0] ou -0.001; origem=1; destino=2", "1. Hypothesis gera valor. 2. POST /transferencia. 3. Validar status.", "HTTP 422 para todo valor nao positivo.", "Sem alteracao de saldo esperada.", "Nenhuma.", "Sistema", "Caixa preta; teste baseado em propriedades e particao de equivalencia.", "Test Design and Execution", "REQ-TRF-002"),
        ("TC-S4-001", "Recusar transferencia quando conta de origem nao existe.", "Flask test client ativo; banco resetado.", "origem=999, destino=1, valor=100.00", "1. Criar client. 2. POST /transferencia. 3. Verificar JSON.", "HTTP 404 e erro Conta nao encontrada.", "Banco permanece sem transferencia registrada.", "Nenhuma.", "Integracao", "Caixa preta; particao de equivalencia para identificador invalido.", "Test Environment", "REQ-TRF-003"),
        ("TC-S5-001", "Aceitar transferencia de pequeno valor positivo.", "Conta 1 com R$ 1000,00 e conta 2 com R$ 500,00.", "origem=2, destino=1, valor=0.50", "1. Given valida saldos. 2. When transfere 0.50. 3. Then valida status e mensagem.", "HTTP 200 e mensagem Transferencia realizada.", "Banco resetado antes do proximo cenario por fixture autouse.", "Contexto do feature.", "Aceitacao", "BDD; caixa preta por regra de negocio.", "Test Design and Execution", "REQ-TRF-004"),
    ]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["ID", "Descricao", "Entradas e passos", "Resultado esperado", "Classificacao e rastreabilidade"]
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
        shade(table.rows[0].cells[idx], LIGHT_GRAY)
        for run in table.rows[0].cells[idx].paragraphs[0].runs:
            set_font(run, bold=True, size=8.6)
    repeat_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row[0]
        cells[1].text = f"{row[1]}\nPre-condicoes: {row[2]}"
        cells[2].text = f"Dados: {row[3]}\nPassos: {row[4]}"
        cells[3].text = f"{row[5]}\nSaida: {row[6]}\nDependencias: {row[7]}"
        cells[4].text = f"Nivel BSTQB: {row[8]}\nTecnica: {row[9]}\nTMMi: {row[10]}\nRequisito: {row[11]}"
        for cell in cells:
            for run in cell.paragraphs[0].runs:
                set_font(run, size=7.9)
    set_table_widths(table, [0.8, 1.45, 1.55, 1.35, 1.35])


def add_schemathesis(doc):
    doc.add_heading("2. OpenAPI 3.1 e Schemathesis", level=1)
    para(doc, "O contrato abaixo descreve os tres endpoints observados da API. O Schemathesis 4.21.10 foi executado com a API local em http://127.0.0.1:5000.")
    doc.add_heading("2.1 Conteudo do openapi.yaml", level=2)
    code_block(doc, read_text(ROOT / "openapi.yaml"), size=6.1)
    doc.add_heading("2.2 Output inicial com achados", level=2)
    code_block(doc, read_text(ROOT / "schemathesis_output_achados.txt"), size=6.4, max_chars=6200)
    doc.add_heading("2.3 Output final estabilizado", level=2)
    code_block(doc, read_text(ROOT / "schemathesis_output_final.txt"), size=6.4)

    doc.add_heading("2.4 Analise dos achados", level=2)
    findings = [
        ("404 HTML em /saldo/null,null e /extrato/null,null", "B - erro de API", "O contrato declarava JSON, mas rotas Flask nao correspondidas retornavam pagina HTML padrao. A correcao foi adicionar errorhandler 404 com JSON."),
        ("POST /transferencia com corpo [null, null] gerou AttributeError", "C - defeito real", "A API assumia que request.get_json retornaria dict. Listas JSON validas para o parser quebravam antes da validacao de campos. A correcao foi rejeitar corpo que nao seja objeto."),
        ("origem igual a destino rejeitada como 422", "A - limite do contrato de teste", "A regra relacional origem != destino e real, mas nao foi expressa no schema simples. Por isso o CI exclui positive_data_acceptance e mantem os checks de erro de servidor, content type, status e schema."),
        ("Warnings finais sobre poucos dados validos em POST /transferencia", "A - erro/limite de teste", "O gerador cria muitas contas inexistentes. O resultado final ainda e util, mas a proxima melhoria seria fornecer exemplos realistas ou seeds para contas 1, 2 e 3."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for idx, h in enumerate(["Achado", "Classificacao", "Justificativa"]):
        table.rows[0].cells[idx].text = h
        shade(table.rows[0].cells[idx], LIGHT_GRAY)
        for run in table.rows[0].cells[idx].paragraphs[0].runs:
            set_font(run, bold=True, size=9)
    for finding in findings:
        cells = table.add_row().cells
        for idx, text in enumerate(finding):
            cells[idx].text = text
            for run in cells[idx].paragraphs[0].runs:
                set_font(run, size=8.8)
    set_table_widths(table, [1.75, 1.25, 3.5])
    para(doc, "Diferenca epistemologica: o Hypothesis do Sprint 3 verificou invariantes formuladas manualmente pelo testador; o Schemathesis verificou se a API obedecia ao contrato declarativo OpenAPI que passou a funcionar como oraculo HTTP.")


def add_ci(doc):
    doc.add_heading("3. Matriz de rastreabilidade automatizada no CI", level=1)
    para(doc, "Foram criados .github/workflows/test.yml e scripts/gera_matriz.py. O workflow instala dependencias, sobe a API, executa pytest com JUnit XML, roda Schemathesis e publica report.xml e matriz_rastreabilidade.md como artefatos.")
    para(doc, "URL do repositorio GitHub: pendente de publicacao. O workspace local nao possui remote Git configurado; portanto, esta entrega registra a execucao local equivalente e os arquivos prontos para push.")
    para(doc, "Print da aba Actions: pendente ate o primeiro push. Evidencia local equivalente: 29 testes pytest passaram, report.xml foi gerado e a matriz abaixo foi produzida automaticamente a partir dele.")
    doc.add_heading("3.1 Conteudo da matriz_rastreabilidade.md", level=2)
    code_block(doc, read_text(ROOT / "matriz_rastreabilidade.md"), size=7.0)
    para(doc, "Qualidade demonstravel mudou de uma saida pontual de terminal para uma evidencia reprodutivel: o mesmo comando gera report.xml e matriz rastreavel, reduzindo a distancia entre execucao tecnica e auditoria.")


def add_report(doc):
    doc.add_heading("4. Relatorio tecnico em dialogo com Nayak et al. (2024)", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A matriz automatizada no CI aumenta a qualidade demonstravel, mas exige governanca contra execucoes improdutivas")
    set_font(r, size=14, bold=True)
    texts = [
        "A automacao da matriz de rastreabilidade desloca a evidencia de qualidade do documento estatico para o pipeline executavel. Minha tese e que essa mudanca aumenta a maturidade do processo, mas tambem cria uma nova responsabilidade: limitar execucoes improdutivas sem reduzir a capacidade detectora da suite.",
        "A evidencia empirica apareceu em tres momentos. Primeiro, a execucao agregada dos testes produziu report.xml com 29 casos verdes, incluindo testes requests, property-based, Flask client e BDD. Segundo, scripts/gera_matriz.py transformou esse XML em matriz requisito-caso-evidencia sem atualizacao manual. Terceiro, o Schemathesis encontrou problemas que a suite manual nao tinha isolado: 404 HTML fora do contrato, corpo JSON em lista gerando 500 e uma regra relacional que o schema nao representava bem.",
        "Nayak et al. (2024) discutem sustentabilidade em continuous testing e chamam atencao para o custo de execucoes nao atendidas em pipelines CI/CD. Concordo com o ponto central: rodar tudo sempre parece rigoroso, mas pode desperdiçar tempo, energia e atencao humana. Minha experiencia acrescenta que a sustentabilidade nao deve ser apenas reduzir runs; ela deve preservar evidencias auditaveis. Um pipeline que economiza execucao mas deixa de gerar matriz perde valor regulatorio.",
        "Proponho uma politica concreta de CI em tres trilhas: em pull request, rodar pytest completo e Schemathesis com max-examples reduzido; em push para main, rodar Schemathesis com mais exemplos e upload obrigatorio da matriz; em agendamento noturno, rodar mutation testing ou contract testing mais agressivo. Alem disso, o workflow deve salvar sempre report.xml, matriz_rastreabilidade.md e output do Schemathesis, para que economia de execucao nao vire perda de evidencia.",
        "O limite do argumento e que a execucao ocorreu localmente e em uma API pequena, sem carga real, sem consumo energetico medido e sem historico de dezenas de commits por dia. Para sustentar a tese em producao, seria necessario publicar o repositorio, coletar tempos de pipeline no GitHub Actions, medir frequencia de falhas, comparar configuracoes de exemplos do Schemathesis e avaliar custo de execucoes completas versus execucoes seletivas.",
    ]
    for text in texts:
        para(doc, text)


def add_synthesis_refs(doc):
    doc.add_heading("5. Sintese epistemologica", level=1)
    para(doc, "a) No Sprint 1, evidencia de qualidade era o pytest verde no terminal; no Sprint 8, ela passou a ser um artefato rastreavel gerado automaticamente a partir do pipeline.")
    para(doc, "b) O achado mais surpreendente do Schemathesis foi o corpo JSON em formato de lista causar erro 500, porque a suite manual testava campos ausentes, mas nao testava tipo estrutural do corpo.")
    para(doc, "c) Minha tese e que qualidade demonstravel exige pipeline, contrato e matriz automatizada, mas tambem exige governanca para evitar execucoes caras que nao aumentam a capacidade detectora.")
    doc.add_heading("6. Referencias", level=1)
    refs = [
        "BSTQB - BRAZILIAN SOFTWARE TESTING QUALIFICATIONS BOARD. Certified Tester Foundation Level Syllabus: versao 4.0. 2023.",
        "DYGALO, D. Schemathesis: Property-based testing for OpenAPI and GraphQL APIs. Documentacao oficial, 2024. Disponivel em: https://schemathesis.readthedocs.io/. Acesso em: jun. 2026.",
        "GITHUB. GitHub Actions Documentation. GitHub, 2024. Disponivel em: https://docs.github.com/en/actions. Acesso em: jun. 2026.",
        "ISO - INTERNATIONAL ORGANIZATION FOR STANDARDIZATION. ISO/IEC/IEEE 29119-3:2021 - Software and systems engineering - Software testing - Part 3: Test documentation. Geneva: ISO, 2021.",
        "NAYAK, K. et al. Sustainable Continuous Testing in DevOps Pipeline. IEEE Conference Publication, 2024.",
        "OPENAPI INITIATIVE. OpenAPI Specification v3.1.0. OpenAPI Initiative / Linux Foundation, 2021. Disponivel em: https://spec.openapis.org/oas/v3.1.0. Acesso em: jun. 2026.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.35)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(ref)
        set_font(r, size=10.5)


def main():
    doc = Document()
    configure(doc)
    add_title(doc)
    add_cases(doc)
    add_schemathesis(doc)
    add_ci(doc)
    add_report(doc)
    add_synthesis_refs(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
