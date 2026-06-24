from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Entrega_Sprint_7_Assessment_TMMi_DORA_IA.docx"
ASSETS = ROOT / "relatorio_assets"
ASSETS.mkdir(exist_ok=True)

BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
BLACK = RGBColor(0, 0, 0)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_GREEN = "E2F0D9"
LIGHT_GOLD = "FFF2CC"


def set_font(run, name="Times New Roman", size=12, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    set_font(run, size=10, color=MUTED)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("Teste de Software | Sprint 7")
    set_font(r, size=10, bold=True, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Pagina ")
    set_font(r, size=10, color=MUTED)
    page_number(footer)


def para(doc, text="", bold_prefix=None, style=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
        r = p.add_run(text)
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run("- ")
    set_font(r)
    r = p.add_run(text)
    set_font(r)
    return p


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SPRINT 7")
    set_font(r, size=13, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Assessment TMMi, Metricas DORA e Position Paper Critico")
    set_font(r, size=24, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Projeto: internet_banking_tests")
    set_font(r, size=13, italic=True, color=MUTED)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Disciplina", "Teste de Software"),
        ("Professor", "Dr. Claudinei de Oliveira"),
        ("Objeto avaliado", "API Flask de Internet Banking e suite de testes dos Sprints 1 a 6"),
        ("Base empirica", "app.py, config.py, conftest.py, test_ia_gerado.py, test_hypothesis.py, test_flask_client.py, features/ e resultados mutmut"),
        ("Data", "24 de junho de 2026"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        shade(row.cells[0], LIGHT_BLUE)
        for run in row.cells[0].paragraphs[0].runs:
            set_font(run, bold=True)
        for run in row.cells[1].paragraphs[0].runs:
            set_font(run)
    set_table_widths(table, [1.65, 4.85])
    doc.add_page_break()


def add_tmmi(doc):
    doc.add_heading("1. Auto-avaliacao TMMi com benchmark financeiro", level=1)
    para(
        doc,
        "A avaliacao abaixo compara a suite local com as areas de processo do TMMi Level 2. "
        "O criterio usado foi evidência empirica observavel: arquivo, teste, configuracao ou resultado medido. "
        "Como referencia externa, Van Veenendaal (2024) relata que o nivel modal no dominio financeiro e o TMMi Level 3 Defined.",
    )

    data = [
        (
            "Test Policy and Strategy",
            "Nao atinge formalmente",
            "Nao ha documento de politica ou estrategia de testes. O arquivo RESPOSTAS_SPRINT5_EDITAVEL.md registra decisoes tecnicas, mas nao estabelece objetivos, escopo, riscos e criterios de saida como politica formal.",
        ),
        (
            "Test Planning",
            "Parcial",
            "Ha planejamento incremental por sprint, evidenciado pelos arquivos de entrega e resultados mutmut. Entretanto, nao ha plano global de testes cobrindo todos os sprints, responsabilidades, cronograma e criterios de aceite.",
        ),
        (
            "Test Monitoring and Control",
            "Atinge parcialmente",
            "Os arquivos RESULTADOS_MUTMUT_SPRINT4.md e RESULTADOS_MUTMUT_SPRINT5.md monitoram 219 mutantes, mortos, sobreviventes e sem teste. Tambem ha registro de 14 testes verdes apos BDD.",
        ),
        (
            "Test Design and Execution",
            "Atinge",
            "A suite aplica teste de API, analise de fronteira para valor zero/negativo, particao de equivalencia para contas existentes/inexistentes, property-based testing em test_hypothesis.py e cenarios BDD em features/transferencia.feature.",
        ),
        (
            "Test Environment",
            "Atinge",
            "config.py externaliza DATABASE e TESTING; conftest.py define fixture autouse para resetar banking.db antes de cada teste e fornece Flask test client. Isso torna o ambiente controlado e reproduzivel.",
        ),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Area TMMi Level 2", "Resultado", "Evidencia empirica"]
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
        shade(table.rows[0].cells[idx], LIGHT_GRAY)
        for run in table.rows[0].cells[idx].paragraphs[0].runs:
            set_font(run, bold=True)
    repeat_header(table.rows[0])
    for area, result, evidence in data:
        cells = table.add_row().cells
        for idx, text in enumerate((area, result, evidence)):
            cells[idx].text = text
            if idx == 1:
                shade(cells[idx], LIGHT_GREEN if result == "Atinge" else LIGHT_GOLD)
            for run in cells[idx].paragraphs[0].runs:
                set_font(run, size=10.5)
    set_table_widths(table, [1.65, 1.25, 3.6])

    para(
        doc,
        "Quanto ao Level 3, ha pouca evidencia de Test Organization, Test Training Program, Non-functional Testing e Peer Reviews. "
        "O projeto possui integracao progressiva entre ciclo de vida e teste, mas ainda como pratica academica individual, nao como processo definido organizacionalmente.",
    )
    para(
        doc,
        "Conclusao tecnica: a suite se posiciona formalmente no TMMi Level 1 Initial com praticas fortes de Level 2, porque nao atende completamente politica e planejamento. "
        "Ela fica abaixo do benchmark TMMi Level 3 Defined identificado por Van Veenendaal (2024) para o setor financeiro, embora ja demonstre evidencias concretas de ambiente, design e monitoramento de testes.",
    )


def build_dora_image():
    rows = [
        ["Metrica DORA", "Definicao operacional", "Valor calculado", "Cluster sugerido"],
        ["Lead Time for Changes", "Tempo da Q8 ate primeiro pytest verde do Sprint 1", "7 dias, estimado", "Medium: dado estimado por ausencia de timestamp original preservado"],
        ["Deployment Frequency", "Sprints entregues por semana no semestre", "6 sprints / 4,86 semanas = 1,23 por semana", "High: cadencia aproximadamente semanal"],
        ["Change Failure Rate", "Mutantes sobreviventes Sprint 4 / total de mutantes", "52 / 219 = 23,7%", "Medium/Low: estabilidade limitada por lacunas detectadas"],
        ["Mean Time to Recovery", "Falha do teste de saldo igual ao valor ate correcao via conftest.py", "7 dias, estimado", "Medium: recuperacao no ciclo seguinte"],
    ]
    width, height = 1900, 760
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_path = next((p for p in ["C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/calibri.ttf"] if Path(p).exists()), None)
    font = ImageFont.truetype(font_path, 30) if font_path else ImageFont.load_default()
    font_bold = ImageFont.truetype(font_path, 31) if font_path else ImageFont.load_default()
    col_widths = [340, 560, 430, 570]
    x_positions = [0]
    for w in col_widths[:-1]:
        x_positions.append(x_positions[-1] + w)
    row_heights = [92, 152, 132, 132, 132]
    y = 0
    for ridx, row in enumerate(rows):
        x = 0
        for cidx, cell in enumerate(row):
            fill = "#E8EEF5" if ridx == 0 else "#FFFFFF"
            draw.rectangle((x, y, x + col_widths[cidx], y + row_heights[ridx]), fill=fill, outline="#8A8A8A", width=2)
            words = cell.split()
            lines = []
            line = ""
            max_chars = max(15, col_widths[cidx] // 17)
            for word in words:
                if len(line) + len(word) + 1 <= max_chars:
                    line = f"{line} {word}".strip()
                else:
                    lines.append(line)
                    line = word
            if line:
                lines.append(line)
            ty = y + 18
            for ln in lines[:5]:
                draw.text((x + 16, ty), ln, font=font_bold if ridx == 0 else font, fill="#111111")
                ty += 34
            x += col_widths[cidx]
        y += row_heights[ridx]
    path = ASSETS / "planilha_dora_sprint7.png"
    image.save(path)
    return path


def add_dora(doc):
    doc.add_heading("2. Calculo numerico das metricas DORA", level=1)
    para(
        doc,
        "Como os arquivos locais foram reconstruidos em 12-06-2026, seus timestamps atuais nao preservam com confiabilidade a data original do primeiro pytest verde da Q8. "
        "Por honestidade metodologica, os dois tempos historicos abaixo sao estimados pelo calendario dos sprints, enquanto os dados de mutacao sao calculados diretamente dos resultados registrados.",
    )
    rows = [
        (
            "Lead Time for Changes",
            "Tempo entre a Q8 de 09-05-2026 e o primeiro pytest verde do test_ia_gerado.py.",
            "7 dias, estimado. Os metadados atuais do arquivo mostram 12-06-2026, indicando que a pasta foi consolidada posteriormente.",
            "Medium. A entrega ocorre no ciclo seguinte, nao em fluxo continuo.",
        ),
        (
            "Deployment Frequency",
            "Quantidade de sprints concluidos por semana ao longo do periodo empirico.",
            "6 sprints entre 09-05-2026 e 12-06-2026: 6 / 4,86 = 1,23 sprints por semana.",
            "High para contexto academico: cadencia proxima de uma entrega semanal.",
        ),
        (
            "Change Failure Rate",
            "Mutantes sobreviventes no Sprint 4 divididos pelo total de mutantes gerados.",
            "52 / 219 = 23,7%. Apos BDD, o valor caiu para 35 / 219 = 16,0%.",
            "Medium/Low. O numero mostra lacunas reais, mas houve melhoria empirica no Sprint 5.",
        ),
        (
            "Mean Time to Recovery",
            "Tempo entre a falha repetida do teste de saldo igual ao valor e a correcao via fixture autouse em conftest.py.",
            "7 dias, estimado, correspondente ao intervalo Sprint 1 -> Sprint 2.",
            "Medium. A recuperacao aconteceu no sprint seguinte, nao no mesmo dia.",
        ),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Metrica", "Definicao operacional", "Valor", "Cluster e justificativa"]
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
        shade(table.rows[0].cells[idx], LIGHT_GRAY)
        for run in table.rows[0].cells[idx].paragraphs[0].runs:
            set_font(run, bold=True, size=9.5)
    repeat_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
            for run in cells[idx].paragraphs[0].runs:
                set_font(run, size=8.7)
    set_table_widths(table, [1.25, 1.85, 1.65, 1.75])

    para(
        doc,
        "Posicionamento geral: se a suite fosse uma equipe real de producao, eu a classificaria como Medium Performer. "
        "A cadencia de entrega e boa, mas a taxa de falha por mutantes sobreviventes e o tempo de recuperacao semanal impedem classificacao High ou Elite.",
    )
    image_path = build_dora_image()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(6.3))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Figura 1 - Planilha DORA exportada como imagem.")
    set_font(r, size=10, italic=True, color=MUTED)


def add_llm_investigation(doc):
    doc.add_heading("3. Confronto com El Haji, Brandt e Zaidman (2024)", level=1)
    para(
        doc,
        "El Haji, Brandt e Zaidman (2024) avaliaram 290 testes gerados por GitHub Copilot em projetos Python e observaram baixa usabilidade: com contexto de suite existente, menos da metade das geracoes passava diretamente; sem contexto, a situacao foi muito pior. Essa conclusao dialoga diretamente com a trajetoria deste projeto.",
    )
    para(
        doc,
        "a) Evidencia dos Sprints 1, 4 e 5: ",
        bold_prefix="Baixa usabilidade confirmada. ",
    )
    bullet(
        doc,
        "No Sprint 1, test_ia_gerado.py usava requests contra servidor externo, dependia de estado persistente do banking.db e tinha teste sensivel a execucao consecutiva: test_transferencia_saldo_igual_valor podia falhar com HTTP 422 por saldo insuficiente se o banco nao fosse resetado.",
    )
    bullet(
        doc,
        "No Sprint 4, 52 de 219 mutantes sobreviveram, indicando que a suite gerada e adaptada ainda nao detectava mudancas sintaticas relevantes.",
    )
    bullet(
        doc,
        "No Sprint 5, os steps BDD precisaram ser revisados para usar Flask test client e a fixture central client, sem criar outro reset de banco nem depender de requests. Esse foi o problema arquitetural central da automacao.",
    )
    para(
        doc,
        "b) Contexto melhora a geracao: quando conftest.py surgiu no Sprint 2, ele formalizou o aprendizado do Sprint 1: reset autouse, app.test_client() e app.config controlado. Isso criou contexto estrutural que permitiu aos testes posteriores imitarem uma arquitetura correta, em vez de repetir o padrao fragil de servidor externo.",
    )
    para(
        doc,
        "c) Copilot como assistente, nao autoridade: o momento decisivo ocorreu quando o teste aparentemente correto de saldo igual ao valor falhou na segunda execucao. A sugestao gerada parecia plausivel, mas ignorava isolamento de estado. A autoridade precisava estar na evidencia empirica do pytest, no reset do banco e no mutation testing, nao na fluencia do codigo gerado.",
    )


def add_position_paper(doc):
    doc.add_heading("4. Position paper critico", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A area Test Environment do TMMi precisa incorporar auditoria de IA para suites LLM-augmented")
    set_font(r, size=14, bold=True)

    paragraphs = [
        "A avaliacao de maturidade de testes ainda trata ambiente, planejamento e execucao como dimensoes majoritariamente humanas e deterministicas. Minha tese e que a area Test Environment do TMMi precisa ser reformulada para incluir uma subpratica explicita de auditoria de testes gerados por IA, pois suites LLM-augmented podem parecer maduras em volume e sintaxe, mas permanecer imaturas em isolamento de estado, arquitetura de execucao e capacidade de matar mutantes.",
        "A evidencia empirica do projeto internet_banking_tests sustenta essa tese. No Sprint 1, test_ia_gerado.py executava chamadas requests contra servidor externo e dependia do estado acumulado do banco; por isso, test_transferencia_saldo_igual_valor podia passar uma vez e falhar depois por saldo insuficiente. No Sprint 2, conftest.py corrigiu a fragilidade com fixture autouse e Flask test client. No Sprint 4, a suite ainda permitiu 52 mutantes sobreviventes entre 219, e no Sprint 5 os cenarios BDD reduziram esse numero para 35 ao transformar lacunas tecnicas em regras de negocio testaveis.",
        "Esses achados convergem com El Haji, Brandt e Zaidman (2024), que mostram baixa usabilidade dos testes Python gerados pelo Copilot e necessidade frequente de modificacao antes do uso direto. Tambem dialogam com Yuan et al. (2024), pois a avaliacao de testes gerados por LLM nao pode parar em compilacao ou aparencia: e preciso verificar comportamento, cobertura, oraculos e adequacao ao contexto. Minha experiencia acrescenta um ponto: mesmo quando o teste gerado passa, ele pode embutir uma premissa operacional errada, como banco persistente sem reset ou cliente HTTP externo em uma suite que deveria usar Flask test client.",
        "Proponho acrescentar ao TMMi uma pratica chamada AI-Augmented Test Audit dentro de Test Environment e Test Design and Execution. Essa pratica exigiria registrar a origem dos testes gerados por IA, revisar dependencias de ambiente, executar os testes duas vezes consecutivas, medir mutation score antes e depois da revisao humana, e justificar cada teste mantido por uma regra de negocio ou propriedade verificavel. No DORA, essa mesma preocupacao poderia aparecer como um AI Test Escape Rate: percentual de falhas, mutantes ou contraexemplos que escaparam de testes gerados por IA e foram detectados apenas por revisao posterior.",
        "O limite do argumento e que o projeto analisado e academico, pequeno e concentrado em uma API Flask com SQLite. Nao ha equipe real, producao real, incidentes reais nem historico completo de commits preservado. Para sustentar a tese com mais robustez, seria necessario comparar varias equipes, registrar prompts, commits, tempos reais de correcao, revisoes por pares e resultados de mutation testing em sistemas financeiros maiores. Ainda assim, o caso mostra que maturidade em testes na era LLM nao pode ser medida apenas por quantidade de testes: precisa medir governanca sobre a origem, o contexto e a capacidade detectora desses testes.",
    ]
    for text in paragraphs:
        para(doc, text)

    doc.add_heading("Referencias do position paper", level=2)
    refs = [
        "DORA / GOOGLE CLOUD. Accelerate State of DevOps Report 2024. Google Cloud, 2024. Disponivel em: https://dora.dev/research/2024/dora-report/. Acesso em: jun. 2026.",
        "EL HAJI, K.; BRANDT, C.; ZAIDMAN, A. Using GitHub Copilot for Test Generation in Python: An Empirical Study. In: ACM/IEEE INTERNATIONAL CONFERENCE ON AUTOMATION OF SOFTWARE TEST, 5., 2024, Lisbon. Proceedings... New York: ACM, 2024. p. 45-55. DOI: 10.1145/3644032.3644443.",
        "LI, Y. et al. Evaluating large language models for software testing. Computer Standards & Interfaces, v. 93, 103942, 2025. DOI: 10.1016/j.csi.2024.103942.",
        "TMMi FOUNDATION. TMMi Framework: Release 1.2. TMMi Foundation, 2018. Disponivel em: https://www.tmmi.org/tmmi-framework/. Acesso em: jun. 2026.",
        "VAN VEENENDAAL, E. Test Maturity Model integration (TMMi): Test Maturity in the Financial Domain. American Journal of Computer Science and Technology, v. 7, n. 2, 2024. DOI: 10.11648/j.ajcst.20240702.13.",
        "YUAN, Z. et al. Evaluating and Improving ChatGPT for Unit Test Generation. Proceedings of the ACM on Software Engineering, v. 1, n. FSE, 2024. DOI: 10.1145/3660783.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.35)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(ref)
        set_font(r, size=10.5)


def add_synthesis(doc):
    doc.add_heading("5. Sintese epistemologica", level=1)
    para(
        doc,
        "a) Saber o nome do que produzi, no Sprint 6, foi classificar artefatos na taxonomia BSTQB; saber o nivel de maturidade, no Sprint 7, foi avaliar se esses artefatos formam um processo sustentavel, medido e comparavel.",
    )
    para(
        doc,
        "b) O numero DORA que mais surpreendeu foi o Change Failure Rate de 23,7%, porque ele mostrou que uma suite com testes verdes ainda deixava quase um quarto dos mutantes escapar no Sprint 4.",
    )
    para(
        doc,
        "c) Minha tese e que modelos de maturidade precisam medir auditoria de IA, pois testes LLM-augmented so amadurecem quando sao confrontados com ambiente controlado, execucao repetida e mutation testing.",
    )


def add_final_refs(doc):
    doc.add_heading("6. Referencias bibliograficas", level=1)
    refs = [
        "DORA / GOOGLE CLOUD. Accelerate State of DevOps Report 2024. Google Cloud, 2024. Disponivel em: https://dora.dev/research/2024/dora-report/. Acesso em: jun. 2026.",
        "EL HAJI, K.; BRANDT, C.; ZAIDMAN, A. Using GitHub Copilot for Test Generation in Python: An Empirical Study. In: ACM/IEEE INTERNATIONAL CONFERENCE ON AUTOMATION OF SOFTWARE TEST, 5., 2024, Lisbon. Proceedings... New York: ACM, 2024. p. 45-55. DOI: 10.1145/3644032.3644443.",
        "LI, Y. et al. Evaluating large language models for software testing. Computer Standards & Interfaces, v. 93, 103942, 2025. DOI: 10.1016/j.csi.2024.103942.",
        "TMMi FOUNDATION. TMMi Framework: Release 1.2. TMMi Foundation, 2018. Disponivel em: https://www.tmmi.org/tmmi-framework/. Acesso em: jun. 2026.",
        "VAN VEENENDAAL, E. Test Maturity Model integration (TMMi): Test Maturity in the Financial Domain. American Journal of Computer Science and Technology, v. 7, n. 2, 2024. DOI: 10.11648/j.ajcst.20240702.13.",
        "YUAN, Z. et al. Evaluating and Improving ChatGPT for Unit Test Generation. Proceedings of the ACM on Software Engineering, v. 1, n. FSE, 2024. DOI: 10.1145/3660783.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.35)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(ref)
        set_font(r, size=10.5)


def main():
    doc = Document()
    configure(doc)
    add_title(doc)
    add_tmmi(doc)
    add_dora(doc)
    add_llm_investigation(doc)
    add_position_paper(doc)
    add_synthesis(doc)
    add_final_refs(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
